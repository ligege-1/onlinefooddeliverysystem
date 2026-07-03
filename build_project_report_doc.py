from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = r"B:\app\Android订餐系统课程项目报告-优化版.docx"

IMAGE_PATHS = {
    "home_top": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-6d9a66d1-3c17-44a9-8504-0d37b1f1ddb2.png",
    "home_bottom": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-6bf39ce0-f036-4501-b574-2b22483d19fb.png",
    "shop_detail": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-9c64d492-4f24-43ee-981b-41cc1f3db88e.png",
    "food_detail": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-3eb151da-62cf-44f8-ac9b-3381e95a4ffd.png",
    "cart_top": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-17f51e00-6271-4d89-89f7-c61fb1866b9c.png",
    "cart_bottom": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-ae2e3a49-152d-4a5f-91d5-44d75be46baf.png",
    "pay_wait": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-ff03b28a-ce61-4a51-9113-b57bb33bc718.png",
    "pay_success": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-e7e8f536-0a0d-40b2-bcd9-573f39a8a0a9.png",
    "orders_mix": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-0c43ea66-6f64-42ff-9c95-1db1f4896192.png",
    "pay_continue": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-71238138-b7f1-4b67-b2c8-ac97ed089163.png",
    "orders_paid": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-c02cf838-964d-4584-846e-22698b87a73e.png",
    "orders_clear_confirm": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-08cd6374-ac1c-428b-903e-b57bb65e8302.png",
    "orders_empty": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-5cc6b0d9-622f-4b52-83cb-86c060e695ae.png",
    "ai_input": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-f903147d-eb04-4a66-a0be-9a22ec2126ea.png",
    "ai_result": r"C:\Users\31349\AppData\Local\Temp\codex-clipboard-0f0a0dfe-997b-4578-939b-a7f5c9c7f9c5.png",
}


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ("val", "sz", "space", "color"):
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_run_font(run, size, bold=False, font_name="SimSun"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0.74, line_spacing=1.5):
    fmt = paragraph.paragraph_format
    fmt.alignment = align
    fmt.first_line_indent = Cm(first_line) if first_line else Cm(0)
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def add_text_paragraph(doc, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0.74, line_spacing=1.5):
    p = doc.add_paragraph()
    set_paragraph(p, align=align, first_line=first_line, line_spacing=line_spacing)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)
    return p


def add_text_paragraph_to_cell(cell, text):
    p = cell.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0.74, line_spacing=1.5)
    run = p.add_run(text)
    set_run_font(run, 12)


def add_title(doc, text, size=22):
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, line_spacing=1.5)
    run = p.add_run(text)
    set_run_font(run, size, bold=True)
    return p


def add_chapter_title(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, line_spacing=1.5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, 16, bold=True)
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0, line_spacing=1.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, 15, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, line_spacing=1.3)
    run = p.add_run(text)
    set_run_font(run, 11)


def add_simple_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        if widths:
            hdr[i].width = widths[i]
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[i].paragraphs[0]
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, line_spacing=1.2)
        for run in p.runs:
            set_run_font(run, 11, bold=True)
    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = text
            if widths:
                row[i].width = widths[i]
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in row[i].paragraphs:
                set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0, line_spacing=1.2)
                for run in p.runs:
                    set_run_font(run, 10.5)
    return table


def add_image(doc, image_path, width_cm=15.5):
    path = Path(image_path)
    if not path.exists():
        add_text_paragraph(doc, f"截图缺失：{image_path}", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0)
        return
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, line_spacing=1.0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def add_code_block(doc, code_lines):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_border(cell,
                    top={"val": "single", "sz": 8, "space": 0, "color": "666666"},
                    left={"val": "single", "sz": 8, "space": 0, "color": "666666"},
                    bottom={"val": "single", "sz": 8, "space": 0, "color": "666666"},
                    right={"val": "single", "sz": 8, "space": 0, "color": "666666"})
    p = cell.paragraphs[0]
    set_paragraph(p, first_line=0, line_spacing=1.15)
    for i, line in enumerate(code_lines):
        run = p.add_run(line)
        set_run_font(run, 10, font_name="Consolas")
        if i < len(code_lines) - 1:
            run.add_break()


def configure_section(section):
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(1.75)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_cover(doc):
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, line_spacing=1.0)
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("成都理工大学")
    set_run_font(run, 28, bold=True, font_name="KaiTi")

    for _ in range(4):
        doc.add_paragraph()

    add_title(doc, "《Android 开发技术》课程项目报告")

    fields = [
        "设计题目    校园订餐系统",
        "学院名称    计算机与网络安全学院",
        "专业名称    物联网工程",
        "学生姓名    李家豪",
        "学生学号    202319120612",
        "任课教师    温泉",
        "设计成绩    __________",
    ]
    for _ in range(2):
        doc.add_paragraph()
    for text in fields:
        p = doc.add_paragraph()
        set_paragraph(p, first_line=0, line_spacing=1.8)
        p.paragraph_format.left_indent = Cm(1.2)
        run = p.add_run(text)
        set_run_font(run, 16, bold=True)

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0)
    run = p.add_run("教务处制")
    set_run_font(run, 14, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0)
    run = p.add_run("2026 年 6 月 12 日")
    set_run_font(run, 14)


def add_manual_toc(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1])
    add_title(doc, "目录")
    entries = [
        "摘要 .................................................... 1",
        "第 1 章 绪论 ............................................. 2",
        "第 2 章 需求分析 ......................................... 4",
        "第 3 章 内容与方法 ....................................... 7",
        "第 4 章 系统详细设计与实现 ............................... 12",
        "学生学习心得 ............................................ 21",
        "任课教师评语 ............................................ 22",
    ]
    for text in entries:
        p = doc.add_paragraph()
        set_paragraph(p, first_line=0, line_spacing=1.8)
        p.paragraph_format.left_indent = Cm(1.2)
        run = p.add_run(text)
        set_run_font(run, 14, bold=text.startswith("第"))


def add_abstract_page(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1])
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0)
    run = p.add_run("题目：校园订餐系统")
    set_run_font(run, 20, bold=True)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(4), Cm(5), Cm(7.5)]
    hdr = table.rows[0].cells
    hdr[0].text = "学生姓名"
    hdr[1].text = "学号"
    hdr[2].text = "完成内容"
    for i, cell in enumerate(hdr):
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, line_spacing=1.2)
        for run in p.runs:
            set_run_font(run, 12, bold=True)

    row = table.add_row().cells
    row[0].text = "李家豪"
    row[1].text = "202319120612"
    row[2].text = "独立完成系统设计、界面开发、购物车与订单功能实现、支付演示、AI 推荐接入以及课程报告撰写"
    for i, cell in enumerate(row):
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT, first_line=0, line_spacing=1.3)
            for run in p.runs:
                set_run_font(run, 12)

    doc.add_paragraph()
    add_title(doc, "摘要")
    add_text_paragraph(
        doc,
        "本项目设计并实现了一个基于 Android 平台的校园在线订餐系统，系统围绕高校学生在校内进行快捷点餐的实际需求展开，"
        "完成了从首页浏览店铺、进入店铺挑选菜品、购物车汇总、订单确认、扫码支付演示到历史订单查看的完整流程。"
        "为了让项目更贴近真实应用，我在界面中使用了真实餐品图片，并对订单确认页、支付页和历史订单页做了多轮优化，"
        "解决了页面信息不够清晰、商品显示不完整以及支付状态不同步等问题。"
        "此外，项目还加入了 AI 智能点餐模块，能够根据用户输入的用餐需求生成店铺和菜品推荐结果，并支持将推荐结果直接加入购物车。"
        "整个系统采用 Java 与 Android 原生组件开发，综合使用了 RecyclerView、Intent、SQLite、本地数据管理和二维码生成等技术，"
        "较好地体现了课程项目对界面实现、业务流程和功能扩展能力的要求。"
    )
    p = doc.add_paragraph()
    set_paragraph(p, first_line=0)
    run = p.add_run("关键词：")
    set_run_font(run, 12, bold=True)
    run = p.add_run("Android；校园订餐；AI 推荐；SQLite；扫码支付演示")
    set_run_font(run, 12)


def add_body(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1])
    add_chapter_title(doc, "第 1 章 绪论")
    add_section_title(doc, "1.1 项目背景")
    add_text_paragraph(
        doc,
        "本项目是《Android 开发技术》课程中的综合实践项目，题目为“校园订餐系统”。"
        "选择订餐系统作为课程项目，一方面是因为该类应用与日常生活联系紧密，页面结构和业务流程都比较典型；"
        "另一方面，这类项目能够把课堂中学到的界面布局、列表展示、数据传递、本地存储和页面跳转等知识点串联起来，"
        "比较适合作为一次完整的课程项目训练。"
    )
    add_text_paragraph(
        doc,
        "本项目并不是简单地绘制几个静态页面，而是希望尽量模拟真实点餐流程。用户进入系统后，"
        "可以先在首页查看店铺，再进入店铺选择菜品，之后完成购物车汇总、订单确认、扫码支付演示以及历史订单查看。"
        "为了让项目内容更完整，我还在基础订餐流程之外加入了 AI 推荐模块，"
        "让系统能够根据用户输入的点餐需求给出推荐结果，提高项目的完整度和展示效果。"
    )
    add_section_title(doc, "1.2 项目目的")
    add_text_paragraph(
        doc,
        "本项目的主要目的有三个。第一，完成一个可运行、可演示、流程完整的 Android 订餐应用；"
        "第二，在项目中综合使用 RecyclerView、Intent、SQLite、图片资源管理等课程知识；"
        "第三，在已有业务功能的基础上增加一定扩展能力，例如 AI 推荐和二维码支付演示，"
        "使项目不仅能完成课程要求，也能体现一定的实际应用感。"
    )
    add_text_paragraph(
        doc,
        "在实现过程中，我除了关注功能能否跑通，也比较重视页面可读性和交互细节。"
        "尤其是订单确认页、支付页和历史订单页这些与演示效果直接相关的部分，"
        "都经过了多轮调整，以保证最终展示时既能看清楚信息，又能完整走通业务流程。"
    )
    add_section_title(doc, "1.3 项目概况")
    add_text_paragraph(
        doc,
        "本系统面向校园日常订餐场景，主要包含首页店铺浏览、店铺详情、菜品详情、购物车与订单确认、扫码支付、历史订单以及 AI 智能点餐几个功能模块。"
        "系统整体采用 Android 原生方式实现，界面层面结合真实餐品图片进行展示，数据层面使用本地 SQLite 保存订单状态。"
        "从最终效果来看，本项目已经能够完成一个较为完整的校园订餐业务演示流程。"
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1])
    add_chapter_title(doc, "第 2 章 需求分析")
    add_section_title(doc, "2.1 功能需求分析")
    add_text_paragraph(
        doc,
        "根据项目内容，系统主要需要完成首页店铺展示、店铺详情展示、菜品加入购物车、订单确认、支付演示、历史订单查看以及 AI 推荐等功能。"
        "首页需要展示多家店铺，并支持搜索；店铺详情页需要展示当前店铺的菜品信息；订单确认页需要汇总用户当前选择的全部商品；"
        "支付页需要支持二维码展示和模拟支付成功；历史订单页需要能够展示订单状态并支持继续支付。"
    )
    add_text_paragraph(
        doc,
        "在这些功能中，比较关键的不是某一个页面能否单独显示出来，而是页面之间的数据是否能够连起来。"
        "例如在店铺页加购商品后，购物车数量要同步；在订单确认页中，所选商品不能漏显示；"
        "支付成功以后，历史订单中的状态也要跟着变化。"
        "因此，项目需求除了“有这些页面”之外，还包括“这些页面之间要形成完整流程”。"
    )
    add_section_title(doc, "2.2 页面设计分析")
    add_text_paragraph(
        doc,
        "在页面设计上，我希望整体效果更接近真实点餐 App，而不是课堂作业里常见的空白页面。"
        "因此在后期优化中，我将默认图片替换成了真实餐品图片，同时调整了列表卡片、按钮、标题栏和信息区的层次。"
        "首页强调店铺卡片的识别性，详情页强调菜品信息完整性，订单确认页则强调信息汇总的清晰度。"
    )
    add_text_paragraph(
        doc,
        "由于订餐系统中图片占比较高，如果图片显示不完整或者裁切太怪，会直接影响页面观感。"
        "所以在商品详情页和订单页中，我对图片展示方式做了调整，尽量让主体内容完整显示，"
        "避免出现商品图只露出一部分、或者重要信息被裁掉的情况。"
    )
    add_text_paragraph(
        doc,
        "另外，页面中还需要明确告诉用户当前操作的是哪一家店、选中了哪些商品、现在处于什么状态。"
        "这也是我后续重点修改的地方。比如在订单确认页中增加店铺信息展示，"
        "在支付页和历史订单页中明确区分“待支付”和“已支付”，这样在演示时会更直观。"
    )
    add_section_title(doc, "2.3 模块划分说明")
    add_text_paragraph(
        doc,
        "为了便于开发和后续修改，我将系统按页面和功能拆分为几个主要模块：首页模块负责店铺展示与搜索，"
        "店铺模块负责菜品列表和详情展示，购物车模块负责数量维护和价格汇总，支付模块负责二维码展示和状态切换，"
        "订单模块负责历史记录和继续支付，AI 模块负责推荐结果生成与跳转加购。"
    )
    add_text_paragraph(
        doc,
        "这种拆分方式的好处是比较清楚：每个页面负责一部分明确任务，出现问题时也比较容易定位。"
        "例如订单状态不对时，就重点检查支付页和订单数据库；图片显示不完整时，就集中修改详情页和列表项布局；"
        "推荐结果无法加入购物车时，就检查推荐模块和购物车管理类之间的数据传递。"
    )
    add_section_title(doc, "2.4 开发环境与技术说明")
    add_text_paragraph(
        doc,
        "从开发环境来看，本项目主要在 Android Studio 中完成，编程语言采用 Java。"
        "界面部分以 Android 原生布局组件为主，列表展示使用 RecyclerView，订单持久化采用 SQLite，"
        "二维码展示结合生成工具完成，AI 推荐部分通过外部服务接口生成推荐结果。"
        "这些技术都是课程学习内容或与课程内容密切相关的扩展点，适合用于综合项目训练。"
    )
    add_simple_table(
        doc,
        ["类别", "名称", "在项目中的作用"],
        [
            ["开发工具", "Android Studio", "用于页面开发、代码编写、运行调试和项目管理"],
            ["开发语言", "Java", "完成 Activity、数据处理、业务逻辑和页面交互实现"],
            ["界面组件", "RecyclerView", "用于首页店铺列表和菜品列表展示"],
            ["本地存储", "SQLite", "保存订单记录、支付状态和历史订单数据"],
            ["二维码工具", "ZXing / 二维码生成工具", "用于支付页生成测试二维码"],
            ["AI 服务", "SiliconFlow 接口", "根据用户输入生成推荐店铺、推荐理由和推荐菜品"],
        ],
        widths=[Cm(3), Cm(4), Cm(9)]
    )
    add_section_title(doc, "2.5 系统模块说明")
    add_text_paragraph(
        doc,
        "为了让系统结构更清楚，也为了让后续答辩或展示时更容易说明，我将主要功能整理为若干独立模块。"
        "每个模块既有自己的页面职责，也会与其它模块发生数据联动，组合后形成完整订餐流程。"
        "下面给出本项目主要模块及其作用说明。"
    )
    add_simple_table(
        doc,
        ["模块名称", "主要页面/类", "功能说明"],
        [
            ["首页模块", "MainActivity", "展示店铺列表、处理关键字搜索、跳转 AI 点餐和历史订单页面"],
            ["店铺模块", "ShopDetailActivity", "展示店铺基本信息、菜品列表并支持加购"],
            ["菜品详情模块", "FoodDetailActivity", "展示单个菜品大图、描述、价格和加入购物车入口"],
            ["购物车模块", "CartActivity / CartManager", "维护商品数量、渲染订单确认页并计算价格"],
            ["支付模块", "PayActivity / QrCodeUtils", "生成测试二维码并模拟支付成功流程"],
            ["订单模块", "OrdersActivity / OrderDbHelper", "保存和展示历史订单，支持继续支付和清空记录"],
            ["AI 模块", "RecommendActivity / SiliconFlowService", "根据用户需求生成推荐套餐并加入购物车"],
        ],
        widths=[Cm(3.2), Cm(4.2), Cm(8.6)]
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1])
    add_chapter_title(doc, "第 3 章 内容与方法")
    add_section_title(doc, "3.1 首页与店铺列表实现")
    add_text_paragraph(
        doc,
        "系统首页由 MainActivity 承担，主要负责展示店铺列表和处理搜索逻辑。为了让首页看起来更像实际应用，"
        "我将店铺信息做成卡片式展示，包括店铺名称、分类、评分、配送信息和推荐语。"
        "用户可以通过搜索框按关键字筛选店铺，也可以直接点击卡片进入店铺详情页。"
    )
    add_text_paragraph(
        doc,
        "在搜索实现上，我没有只按店铺名做最简单的匹配，而是尽量让搜索更实用。"
        "除了店铺名称外，也可以根据分类和部分菜品信息进行筛选，这样用户在演示时更容易看出系统具备一定交互能力。"
    )
    add_section_title(doc, "3.2 菜品详情与购物车实现")
    add_text_paragraph(
        doc,
        "进入店铺详情页后，系统会展示当前店铺的菜品列表。每个菜品项都包含图片、名称、价格、描述和加减按钮，"
        "用户既可以直接在列表中调整数量，也可以进入详情页查看大图和更完整的信息。"
        "这一部分的重点在于让商品展示更真实，所以我对图片资源进行了替换，尽量使用更贴近外卖平台视觉效果的餐品图片。"
    )
    add_text_paragraph(
        doc,
        "购物车部分由统一的管理类维护当前已选商品及其数量。这样做的原因是，用户可能在多个页面里对商品数量进行修改，"
        "如果没有统一管理，页面之间的数据很容易不同步。通过集中维护购物车状态，能够让店铺页、详情页和订单确认页看到的是同一份数据。"
    )
    add_text_paragraph(
        doc,
        "订单确认页是系统里改动最大的一部分之一。起初页面存在内容偏空、店铺信息不够明显、商品显示不完整等问题，"
        "后面我重新整理了页面结构，把它拆成地址区、店铺信息区、商品列表区和费用明细区。"
        "这样处理之后，页面信息层次更清楚，用户也能一眼看出自己是在给哪家店下单。"
    )
    add_section_title(doc, "3.3 支付、订单与 AI 推荐实现")
    add_text_paragraph(
        doc,
        "支付页采用二维码演示的方式来模拟真实支付场景。考虑到课堂项目无法真正接入线上支付回调，"
        "我保留了扫码支付的形式，同时增加“我已扫码，模拟支付成功”的按钮，让整个演示流程既有实际感，也能够稳定完成。"
    )
    add_text_paragraph(
        doc,
        "用户点击模拟支付成功后，系统不仅会在当前页面更新为“支付成功”状态，还会把订单状态写回本地存储。"
        "这样做的目的，是保证历史订单页中的记录和支付结果保持一致，避免出现当前页显示支付成功、历史页却仍然显示待支付的问题。"
    )
    add_text_paragraph(
        doc,
        "AI 推荐模块是本项目中比较有特色的一部分。用户可以输入自己的需求，例如想吃辣、预算多少、几个人用餐等，"
        "系统再结合当前已有菜单调用 AI 服务生成推荐结果。推荐结果中不仅会给出店铺和菜品，还会说明推荐理由，"
        "并支持把推荐商品加入购物车。这样一来，AI 功能就不只是展示文字，而是和系统的实际下单流程连接起来了。"
    )
    add_section_title(doc, "3.4 数据结构与订单存储方法")
    add_text_paragraph(
        doc,
        "本项目虽然没有接入完整后端，但为了让业务流程更完整，仍然设计了较清晰的数据结构。"
        "其中，ShopBean 用于描述店铺信息，FoodBean 用于描述菜品信息，CartItem 用于表示购物车中的单项商品，"
        "OrderBean 用于保存一次订单的完整信息。通过这些模型类，页面之间在传递数据时会更清楚，也更便于后续扩展。"
    )
    add_text_paragraph(
        doc,
        "在订单存储方面，系统采用 SQLite 作为本地持久化方案。"
        "这样做的优点是实现简单、便于课程项目落地，而且能够满足历史订单展示、状态修改和继续支付等需求。"
        "支付成功后，系统会把订单状态更新为“已支付”；如果未支付，则历史订单页中仍然保持“待支付”状态，"
        "用户可以再次进入支付页继续完成本次支付流程。"
    )
    add_simple_table(
        doc,
        ["数据对象", "主要字段", "作用说明"],
        [
            ["ShopBean", "店铺名、分类、评分、配送费、配送时间", "用于首页和店铺详情页展示店铺基础信息"],
            ["FoodBean", "菜品名、价格、口味、描述、图片资源", "用于菜品列表和菜品详情页展示"],
            ["CartItem", "菜品对象、数量、小计", "用于购物车和订单确认页汇总当前所选商品"],
            ["OrderBean", "店铺名、地址、总价、配送费、状态、创建时间", "用于保存完整订单信息并写入本地数据库"],
        ],
        widths=[Cm(3), Cm(6), Cm(7.2)]
    )
    add_text_paragraph(
        doc,
        "这种以模型类加本地数据库的实现方式，虽然比不上完整的前后端分离架构，但对于课程项目来说已经能够很好支撑业务流程。"
        "同时，它也为后续继续扩展真实用户系统、在线订单同步和后端接口预留了空间。"
    )
    add_section_title(doc, "3.5 关键代码说明")
    add_text_paragraph(
        doc,
        "本项目并不需要像论文那样进行大篇幅代码分析，但为了说明项目确实完成了核心功能实现，"
        "这里选取了几段比较关键的代码进行简要说明。每段代码都对应一个实际功能点，重点说明其作用和实现思路。"
    )
    add_section_title(doc, "3.5.1 首页搜索与店铺刷新")
    add_code_block(doc, [
        "etSearch.addTextChangedListener(new TextWatcher() {",
        "    @Override",
        "    public void onTextChanged(CharSequence s, int start, int before, int count) {",
        "        refreshShops(s.toString());",
        "    }",
        "});",
        "",
        "private void refreshShops(String keyword) {",
        "    displayShops.clear();",
        "    displayShops.addAll(DataRepository.searchShops(keyword));",
        "    adapter.notifyDataSetChanged();",
        "}",
    ])
    add_text_paragraph(
        doc,
        "这段代码来自首页 MainActivity，核心作用是监听搜索框内容变化，并在用户输入时实时刷新店铺列表。"
        "通过调用 DataRepository.searchShops(keyword)，首页可以根据关键字重新筛选并展示店铺数据。"
    )
    add_section_title(doc, "3.5.2 购物车渲染与价格汇总")
    add_code_block(doc, [
        "private void refreshCart() {",
        "    cartItems.clear();",
        "    cartItems.addAll(CartManager.getInstance().getItems());",
        "    renderCartItems();",
        "",
        "    int goodsTotal = CartManager.getInstance().getTotalPrice();",
        "    int total = goodsTotal + shop.getDeliveryFee();",
        "    tvGoodsTotal.setText(FormatUtils.price(goodsTotal));",
        "    tvTotal.setText(\"合计 \" + FormatUtils.price(total));",
        "}",
    ])
    add_text_paragraph(
        doc,
        "这段代码来自 CartActivity，用于在订单确认页刷新购物车内容并重新计算价格。"
        "系统先从 CartManager 中读取当前已选商品，再渲染所有商品项，最后根据商品总价和配送费计算订单合计。"
    )
    add_section_title(doc, "3.5.3 支付成功后的订单状态回写")
    add_code_block(doc, [
        "private void showPaidState(boolean saveStatus) {",
        "    if (saveStatus && order != null && order.getId() > 0) {",
        "        new OrderDbHelper(this).updateStatus(order.getId(), \"已支付\");",
        "    }",
        "    tvTitle.setText(\"支付成功\");",
        "    tvHint.setText(\"测试支付已完成，可以查看历史订单\");",
        "    flQrBox.setVisibility(View.GONE);",
        "    tvMockSuccess.setVisibility(View.GONE);",
        "}",
    ])
    add_text_paragraph(
        doc,
        "这段代码来自 PayActivity，作用是在用户点击“模拟支付成功”后更新当前订单状态。"
        "这里通过 OrderDbHelper.updateStatus(...) 将状态真正写回本地数据库，这样历史订单页重新读取数据时就会显示“已支付”。"
    )
    add_section_title(doc, "3.5.4 AI 推荐结果生成与展示")
    add_code_block(doc, [
        "aiService.recommend(prompt, new SiliconFlowService.Callback() {",
        "    @Override",
        "    public void onSuccess(AiRecommendationResult result) {",
        "        tvRun.setEnabled(true);",
        "        tvRun.setText(\"生成推荐\");",
        "        showResult(result);",
        "    }",
        "",
        "    @Override",
        "    public void onError(String message) {",
        "        tvRun.setEnabled(true);",
        "        tvRun.setText(\"生成推荐\");",
        "        tvResult.setText(message);",
        "    }",
        "});",
    ])
    add_text_paragraph(
        doc,
        "这段代码来自 RecommendActivity，负责调用 AI 服务并根据返回结果刷新页面。"
        "当推荐成功时，系统会展示推荐店铺、推荐菜品和推荐理由；当接口返回错误时，则直接把错误信息显示给用户。"
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1])
    add_chapter_title(doc, "第 4 章 系统详细设计与实现")
    add_section_title(doc, "4.1 关键问题与优化过程")
    add_text_paragraph(
        doc,
        "在项目后期调整时，我发现最明显的问题不是功能缺失，而是有些页面虽然能用，但视觉上比较空，尤其是订单确认页，"
        "用户很难一眼看清当前订单属于哪家店。这个问题如果不处理，在演示时会让系统显得不像一个完整应用。"
    )
    add_text_paragraph(
        doc,
        "针对这个问题，我增加了更明确的店铺信息展示区域，并重新组织了结算页的布局，"
        "把原本比较松散的信息集中到几个清晰的模块中。修改之后，用户能够较快看清当前店铺、商品内容和费用明细，页面整体也不再显得空。"
    )
    add_section_title(doc, "4.2 商品显示与结算优化")
    add_text_paragraph(
        doc,
        "另一个比较直接的问题是商品显示不完整。实际测试中出现过选了三个商品、结算区却只看到两个的情况，"
        "这会直接影响订单确认的准确性。对用户来说，结算页最重要的就是看清楚自己到底买了什么，因此这个问题必须优先处理。"
    )
    add_text_paragraph(
        doc,
        "后续我对结算区的商品渲染方式进行了调整，确保所有选中的商品都能完整显示出来。"
        "同时，也让数量变化、单项价格和总价汇总保持一致。这样一来，订单确认页不仅更完整，也更符合真实下单场景。"
    )
    add_section_title(doc, "4.3 支付与订单状态处理")
    add_text_paragraph(
        doc,
        "支付部分最开始也出现过一个比较明显的问题：当前页面提示支付成功，但历史订单中仍然显示待支付。"
        "这个问题本质上是状态没有真正写回去，只是界面上做了一个临时变化。如果不处理，整个支付流程就会显得不可信。"
    )
    add_text_paragraph(
        doc,
        "针对这个问题，我在支付成功后增加了订单状态回写逻辑，并让历史订单页在重新进入时读取最新状态。"
        "这样当前页面和历史页面可以保持一致，系统行为也更符合用户预期。另外，我还保留了继续支付功能，确保待支付订单可以再次进入支付页完成流程。"
    )
    add_section_title(doc, "4.4 图片显示效果优化")
    add_text_paragraph(
        doc,
        "项目中大量使用了餐品图片，因此图片显示效果会直接影响页面质感。如果图片裁切不合理，"
        "就会出现主体不完整、比例奇怪或者看起来很别扭的问题。这一点在商品详情页和订单页里表现得尤其明显。"
    )
    add_text_paragraph(
        doc,
        "为此，我对图片展示策略做了调整，尽量让重要主体显示完整，同时保证不同页面中的图片风格保持统一。"
        "经过这一轮修改后，页面的视觉效果比最初版本更自然，也更适合用于课程演示。"
    )
    add_section_title(doc, "4.5 测试说明")
    add_text_paragraph(
        doc,
        "在测试过程中，我主要从几个方面进行检查。第一是页面展示是否正常，包括首页、详情页、结算页和支付页中的布局、文字和图片；"
        "第二是数据联动是否正常，例如加购后数量是否变化、总价是否正确、订单状态是否同步；"
        "第三是流程是否完整，即从选择商品到支付成功再到历史订单查看，是否能够连续完成。"
    )
    add_text_paragraph(
        doc,
        "结合实际演示情况来看，系统目前已经能够较稳定地完成主要业务流程。AI 推荐模块能够生成结果，"
        "推荐套餐可以加入购物车；支付页可以展示测试二维码；订单状态在历史页中也能够正确反映。"
        "对于课程项目来说，这一版本已经能够较好体现 Android 应用开发中的页面组织、数据管理和功能实现能力。"
    )
    add_section_title(doc, "4.6 系统运行结果展示")
    screenshot_sections = [
        ("4.6.1 首页店铺列表展示", "首页展示了搜索框、AI 点餐入口、历史订单入口以及多家店铺卡片，方便用户从一开始就完成浏览和筛选。", [
            ("图 4-1 首页店铺列表（上半部分）", IMAGE_PATHS["home_top"]),
            ("图 4-2 首页店铺列表（下半部分）", IMAGE_PATHS["home_bottom"]),
        ]),
        ("4.6.2 店铺详情与菜品列表展示", "店铺详情页可以展示店铺基本信息、推荐文案和菜品列表，用户可以直接在列表中进行加购。", [
            ("图 4-3 店铺详情页", IMAGE_PATHS["shop_detail"]),
        ]),
        ("4.6.3 菜品详情页展示", "菜品详情页主要用于查看单个商品的大图、介绍、销量和加入购物车入口。", [
            ("图 4-4 菜品详情页", IMAGE_PATHS["food_detail"]),
        ]),
        ("4.6.4 订单确认页展示", "订单确认页分为地址区、店铺区、商品区和费用区，两张截图组合后可以完整反映本次下单内容。", [
            ("图 4-5 订单确认页（上半部分）", IMAGE_PATHS["cart_top"]),
            ("图 4-6 订单确认页（下半部分）", IMAGE_PATHS["cart_bottom"]),
        ]),
        ("4.6.5 扫码支付与支付成功展示", "支付页提供测试二维码和模拟支付成功按钮，支付完成后页面会切换为支付成功状态。", [
            ("图 4-7 扫码支付页", IMAGE_PATHS["pay_wait"]),
            ("图 4-8 支付成功页", IMAGE_PATHS["pay_success"]),
        ]),
        ("4.6.6 历史订单与状态切换展示", "历史订单页可以清楚区分待支付和已支付状态，并支持从待支付订单继续进入支付页。", [
            ("图 4-9 历史订单页（待支付与已支付同时展示）", IMAGE_PATHS["orders_mix"]),
            ("图 4-10 待支付订单继续支付页", IMAGE_PATHS["pay_continue"]),
            ("图 4-11 历史订单页（支付完成后）", IMAGE_PATHS["orders_paid"]),
        ]),
        ("4.6.7 历史订单清空功能展示", "为了方便测试和演示，系统增加了历史订单清空功能，并在清空前给出确认提示。", [
            ("图 4-12 清空历史订单确认弹窗", IMAGE_PATHS["orders_clear_confirm"]),
            ("图 4-13 历史订单为空时的页面展示", IMAGE_PATHS["orders_empty"]),
        ]),
        ("4.6.8 AI 智能点餐展示", "AI 点餐页面支持输入自然语言需求，并在生成推荐后展示店铺、推荐理由、推荐菜品和预计合计。", [
            ("图 4-14 AI 点餐输入页", IMAGE_PATHS["ai_input"]),
            ("图 4-15 AI 点餐推荐结果页", IMAGE_PATHS["ai_result"]),
        ]),
    ]
    for title, desc, images in screenshot_sections:
        add_section_title(doc, title)
        add_text_paragraph(doc, desc)
        for caption, image_path in images:
            add_image(doc, image_path)
            add_caption(doc, caption)


def add_reflection(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1])
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.width = Cm(2.2)
    right.width = Cm(13.8)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_border(left, top={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    left={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    right={"val": "single", "sz": 8, "space": 0, "color": "000000"})
    set_cell_border(right, top={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    left={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    right={"val": "single", "sz": 8, "space": 0, "color": "000000"})

    p = left.paragraphs[0]
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, line_spacing=2.0)
    for text in ["学", "生", "学", "习", "心", "得"]:
        r = p.add_run(text + "\n")
        set_run_font(r, 14, bold=True)

    add_text_paragraph_to_cell(
        right,
        "通过本次课程项目的独立完成，我对 Android 应用开发有了更加系统和深入的认识。项目并不是简单地将多个页面拼接在一起，"
        "而是要围绕真实业务流程进行设计，从用户进入首页开始，到浏览店铺、查看菜品、加入购物车、确认订单、支付演示，再到历史订单查看，"
        "每一个环节都需要考虑页面结构是否合理、信息展示是否清晰以及状态变化是否同步。"
    )
    add_text_paragraph_to_cell(
        right,
        "在开发过程中，我不仅巩固了 RecyclerView、Activity 跳转、Intent 传值、本地 SQLite 存储等基础知识，"
        "还进一步掌握了图片资源替换、二维码生成、页面样式优化和订单状态回写等更偏综合性的开发技巧。"
        "尤其是在处理订单确认页只显示部分商品、支付成功后历史订单状态未更新等问题时，"
        "我更加深刻地体会到调试、定位问题和逐步修正的重要性。"
    )
    add_text_paragraph_to_cell(
        right,
        "本项目中接入 AI 推荐功能也让我对“传统移动应用 + 智能服务”的结合方式有了初步实践经验。"
        "通过将用户自然语言需求与已有菜单数据一同提供给模型，再将返回结果组织成推荐页面，"
        "我认识到 AI 能力不只是额外加分项，它还可以成为实际业务流程中的一部分。"
    )
    add_text_paragraph_to_cell(
        right,
        "整体来看，这次综合项目让我从单个控件和单个页面的实现，提升到了对完整应用结构的把握。"
        "如果后续继续深入，我希望能够将本项目扩展为带真实后端接口的版本，完善登录、订单同步、真实支付回调等功能，"
        "从而让整个系统更加接近实际可用的产品形态。"
    )
    p = right.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=0)
    run = p.add_run("\n学生签名：__________\n2026 年 6 月")
    set_run_font(run, 12)


def add_teacher_page(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1])
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.width = Cm(2.2)
    right.width = Cm(13.8)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(left, top={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    left={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    right={"val": "single", "sz": 8, "space": 0, "color": "000000"})
    set_cell_border(right, top={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    left={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
                    right={"val": "single", "sz": 8, "space": 0, "color": "000000"})

    p = left.paragraphs[0]
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, line_spacing=2.0)
    for text in ["任", "课", "教", "师", "评", "语"]:
        r = p.add_run(text + "\n")
        set_run_font(r, 14, bold=True)

    p = right.paragraphs[0]
    set_paragraph(p, first_line=0, line_spacing=1.8)
    run = p.add_run("\n\n成绩评定：__________\n\n任课教师（签名）：__________\n\n2026 年 7 月")
    set_run_font(run, 12)


def main():
    doc = Document()
    configure_section(doc.sections[0])
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.font.size = Pt(12)

    add_cover(doc)
    add_manual_toc(doc)
    add_abstract_page(doc)
    add_body(doc)
    add_reflection(doc)
    add_teacher_page(doc)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
